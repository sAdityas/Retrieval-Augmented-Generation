from docx import Document

# Create DOCX file
file_path = "SAP_TCodes_Master_Template_Filled.docx"
doc = Document()
doc.add_heading("SAP T-Codes Master Template", 0)
doc.add_paragraph("This document contains standard SAP T-Codes grouped by module, with placeholders for Description, Cause, and Step-by-Step Solution.")

def add_module(module_name, tcode_list):
    doc.add_page_break()
    doc.add_heading(module_name, level=0)
    for tcode in tcode_list:
        doc.add_heading(tcode, level=1)
        doc.add_paragraph("Description: [Add description here]")
        doc.add_paragraph("Common Cause: [Add cause here]")
        doc.add_paragraph("Steps to Solve: [Add steps here]")
        doc.add_paragraph("\n")

# SAP modules and sample T-Codes
sap_modules = {
    "FI - Financial Accounting": ["FS00", "FB50", "F-02", "FBL3N", "F-28"],
    "CO - Controlling": ["KP06", "KSB1", "KKAO", "KSU5"],
    "MM - Materials Management": ["ME21N", "ME51N", "MIGO", "MB1A", "MB5T"],
    "SD - Sales & Distribution": ["VA01", "VA02", "VL01N", "VF01", "VK11"],
    "PP - Production Planning": ["CO01", "MF50", "MD01", "MD04"],
    "HR - Human Resources": ["PA30", "PT60", "PA20"],
    "PM - Plant Maintenance": ["IW31", "IW32", "IW28"],
    "QM - Quality Management": ["QA32", "QE01", "QA11"],
    "WM - Warehouse Management": ["LT01", "LT10", "LX02"],
    "PS - Project System": ["CJ20N", "CJ30", "CN41N"],
    "CS - Customer Service": ["IW21", "IW22", "IW23"],
    "BI - Business Intelligence": ["RSA1", "RSRT", "RSECADMIN"]
}

for module_name, tcode_list in sap_modules.items():
    add_module(module_name, tcode_list)

# Save final DOCX
doc.save(file_path)
